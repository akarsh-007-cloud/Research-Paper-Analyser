# output_generator.py
import json
from docx import Document

def save_as_json(data: dict, output_path: str):
    """Saves the dictionary data as a JSON file."""
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        print(f"Successfully saved analysis to {output_path}")
    except Exception as e:
        print(f"Error saving to JSON: {e}")

def save_as_docx(data: dict, output_path: str):
    """Saves the dictionary data as a formatted DOCX file."""
    try:
        doc = Document()
        doc.add_heading(data.get("title", "No Title Found"), level=1)

        doc.add_heading("Authors", level=2)
        authors = data.get("authors", ["N/A"])
        doc.add_paragraph(", ".join(authors))

        doc.add_heading("Publication Information", level=2)
        doc.add_paragraph(data.get("publication_info", "N/A"))

        doc.add_heading("Abstract", level=2)
        doc.add_paragraph(data.get("abstract", "N/A"))

        doc.add_heading("Key Findings", level=2)
        # Handle key findings which might be a single string with newlines
        key_findings = data.get("key_findings", "N/A").split('\n')
        for finding in key_findings:
            if finding.strip(): # Avoid adding empty bullet points
                doc.add_paragraph(finding.strip(), style='List Bullet')
        
        doc.save(output_path)
        print(f"Successfully saved analysis to {output_path}")
    except Exception as e:
        print(f"Error saving to DOCX: {e}")